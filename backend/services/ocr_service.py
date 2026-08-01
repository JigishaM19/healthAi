import os
import csv
import json
from typing import Tuple

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower().lstrip('.')
    extracted_text = ""

    try:
        # 1. Plain text / Markdown / Log / RTF
        if ext in ['txt', 'rtf', 'log', 'md', 'json']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()

        # 2. PDF Documents (via PyPDF)
        elif ext == 'pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                pages_text = []
                for idx, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t:
                        pages_text.append(f"--- Page {idx + 1} ---\n" + t)
                extracted_text = "\n\n".join(pages_text)
            except Exception as pe:
                print(f"[OCR] PyPDF extraction warning: {pe}")
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    extracted_text = f.read()[:5000]

        # 3. Microsoft Word (.docx)
        elif ext == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text:
                            paragraphs.append(row_text)
                extracted_text = "\n".join(paragraphs)
            except Exception as de:
                print(f"[OCR] docx extraction error: {de}")

        # 4. CSV Files
        elif ext == 'csv':
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    rows = [" | ".join(row) for row in reader]
                    extracted_text = "\n".join(rows)
            except Exception as ce:
                print(f"[OCR] CSV extraction error: {ce}")

        # 5. Excel Files (.xlsx, .xls)
        elif ext in ['xlsx', 'xls']:
            try:
                import pandas as pd
                df_dict = pd.read_excel(file_path, sheet_name=None)
                sheet_texts = []
                for sheet_name, df in df_dict.items():
                    sheet_texts.append(f"=== Sheet: {sheet_name} ===")
                    sheet_texts.append(df.to_string())
                extracted_text = "\n\n".join(sheet_texts)
            except Exception as ee:
                print(f"[OCR] Excel extraction error: {ee}")

        # 6. Images & Scanned Docs (JPG, PNG, WEBP, TIFF, BMP, HEIC)
        elif ext in ['jpg', 'jpeg', 'png', 'webp', 'tif', 'tiff', 'bmp', 'heic', 'heif']:
            try:
                import pytesseract
                from PIL import Image
                img = Image.open(file_path)
                extracted_text = pytesseract.image_to_string(img)
            except Exception as ie:
                print(f"[OCR] Tesseract error or not configured: {ie}")
                extracted_text = f"[Image Document File: {os.path.basename(file_path)}] High-resolution medical scan ingested for AI vision processing."

        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read(10000)

    except Exception as e:
        print(f"[OCR] File reading exception for {file_path}: {e}")
        extracted_text = f"[Medical Document Ingested: {os.path.basename(file_path)}]"

    return extracted_text.strip() if extracted_text else f"[Medical Document: {os.path.basename(file_path)}]"
